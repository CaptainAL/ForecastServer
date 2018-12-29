# -*- coding: utf-8 -*-
"""
Created on Tue Oct 11 17:44:10 2016

@author: alex.messina
"""



if __name__=='__main__':
    import urllib
    from docx import *
    from docx.shared import Inches
    import datetime as dt
    import time    
    
    from email.MIMEMultipart import MIMEMultipart
    from email.MIMEText import MIMEText
    from email.MIMEImage import MIMEImage
    import smtplib
    
    from PIL import Image
    from PIL import ImageFont
    from PIL import ImageDraw 
    
    def Save_Forecast(site_name,url_list,email=False):
        print 'Looking up weather forecast for '+site_name+'....'
        lat,lon = url_list[0].split('lat=')[1][:7], url_list[0].split('lon=')[1][:9]
        print 'lat/lon: '+str(lat)+','+str(lon)
        print
        maindir = 'C:/Users/alex.messina/Documents/GitHub/ForecastServer/'+site_name + '/'
        ##  Create Document
        document = Document()
        ## Create Email
        msg = MIMEMultipart()
        ## Today's date
        now = str(dt.datetime.now())[:-13]+'00'
        ## Add  Heading to Document
        header = 'NWS Forecast for '+site_name+' - '+now
        document.add_heading(header,level=1)
        document.add_heading('lat/lon: '+str(lat)+','+str(lon),level=1)
        ## Add Heading to email
        msg['SUBJECT']=header
        msg.attach(MIMEText(header+ '    lat/lon: '+str(lat)+','+str(lon)))
        ##
        num_count = 1
        for url in url_list:
            #print url
            ## Get and Save Picture
            pic_filename = maindir+'NWS '+site_name+' Forecast'+str(num_count)+'.png'
            urllib.urlretrieve(url, pic_filename)
            
            
            ## Add location to picture          
            font = ImageFont.truetype("arial.ttf", 20)
            img = Image.open(pic_filename)
            text_box = Image.new("RGB",(150,40),color='black')
            text_draw = ImageDraw.Draw(text_box)
            text_draw.text((10,10),site_name,font=font,fill='white')
            img.paste(text_box)
            img.save(pic_filename)

            ## Insert Picture
            document.add_picture(pic_filename ,width=Inches(7)) ## add pic from filename defined above
            ## Add Picture to email
            msg.attach(MIMEImage(file(pic_filename,'rb').read()))
            ## count Up
            num_count += 1
            
        ## Save document
        print 'Saving forecasts...'
        document.save(maindir+'NWS Forecast for '+site_name+' - '+now+'.docx')
        if email==True:
            ## Send Email
            # to send
            mailer = smtplib.SMTP_SSL('smtp.googlemail.com:465')
            username, password = 'sebal.py', 'Mactec101'
            mailer.login(username,password)
            mailer.sendmail('sebal.py@gmail.com','atm1984@gmail.com', msg.as_string())
            mailer.close()
        print
        print 'Forecasts saved uh successfurry . Have a nice day :)'
        print
        return
    
    def site_url_list(lat,lon,wfo,zcode,gset):
        url_list = ['http://forecast.weather.gov/meteograms/Plotter.php?lat='+lat+'&lon='+lon+'&wfo='+wfo+'&zcode='+zcode+'&gset='+str(gset)+'&gdiff=3&unit=0&tinfo=PY8&ahour=0&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat='+lat+'&lon='+lon+'&wfo='+wfo+'&zcode='+zcode+'&gset='+str(gset)+'&gdiff=3&unit=0&tinfo=PY8&ahour=48&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat='+lat+'&lon='+lon+'&wfo='+wfo+'&zcode='+zcode+'&gset='+str(gset)+'&gdiff=3&unit=0&tinfo=PY8&ahour=96&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6']
        return url_list

      
    ## San Diego (Airport, KSAN)
    lat, lon = '32.7157',  '-117.1617'
    wfo, zcode, gset = 'SGX', 'CAZ043', 18
    SanDiego_url_list = site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Laguna Beach
    lat, lon = '33.5422', '117.7831'
    wfo, zcode, gset = 'SGX', 'CAZ552', 18
    LagunaBeach_url_list = site_url_list(lat,lon,wfo,zcode,gset)
    
    ## March AFB/Riverside
    lat, lon = '33.899', '-117.2522'
    wfo, zcode, gset = 'SGX', 'CAZ048', 18
    Riverside_url_list = site_url_list(lat,lon,wfo,zcode,gset)

    ## Avalon, Catalina Island
    lat, lon = '33.3428', '-118.3278'
    wfo, zcode, gset = 'LOX', 'CAZ087', 15
    Avalon_url_list = site_url_list(lat,lon,wfo,zcode,gset)

    ## Dana Point
    lat, lon = '33.4805', '-117.6977'
    wfo, zcode, gset = 'LOX', 'CAZ087', 15
    DanaPoint_url_list=site_url_list(lat,lon,wfo,zcode,gset)
    
    ## John Wayne Airport
    lat, lon = '33.688', '-117.9043'
    wfo, zcode, gset = 'LOX', 'CAZ087', 15
    JohnWayne_url_list= site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Foothill Ranch
    lat, lon = '33.688', '-117.9043'
    wfo, zcode, gset = 'LOX', 'CAZ087', 15
    FoothillRanch_url_list= site_url_list(lat,lon,wfo,zcode,gset)

    ## Imperial Beach
    lat, lon = '32.5766', '-117.1164'
    wfo, zcode, gset = 'SGX', 'CAZ043', 18
    ImperialBeach_url_list = site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Irvine
    lat, lon = '32.6412', '-117.860'
    wfo, zcode, gset = 'SGX', 'CAZ552', 18
    Irvine_url_list= site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Chollas
    lat, lon = '32.7147', '-117.1241'
    wfo, zcode, gset = 'SGX', 'CAZ043', 18
    Chollas_url_list= site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Lakeside
    lat, lon = '32.854', '-116.9045'
    wfo, zcode, gset = 'SGX', 'CAZ050', 18
    Lakeside_url_list= site_url_list(lat,lon,wfo,zcode,gset)

    
    ## Riding Park (SJC)
    lat, lon = '33.4979', '-117.6654'
    wfo, zcode, gset = 'SGX', 'CAZ552', 18
    RidingPark_url_list= site_url_list(lat,lon,wfo,zcode,gset)    
    
    ## SR73
    lat, lon = '33.6873', '-117.8259'
    wfo, zcode, gset = 'SGX', 'CAZ554', 18
    SR73_url_list= site_url_list(lat,lon,wfo,zcode,gset)    
    
    ## Fallbrook
    lat, lon = '33.3847', '-117.2532'
    wfo, zcode, gset = 'SGX', 'CAZ050', 18
    Fallbrook_url_list = site_url_list(lat,lon,wfo,zcode,gset)
    
    ## Camp Pendleton
    lat, lon = '33.4003', '-117.5765'
    wfo, zcode, gset = 'SGX', 'CAZ043', 18
    Pendleton_url_list = site_url_list(lat,lon,wfo,zcode,gset)

    ## Santa Ana Fire
    lat, lon = '33.7599', '-117.8442'
    wfo, zcode, gset = 'SGX', 'CAZ554', 18
    SantaAnaFire_url_list = site_url_list(lat,lon,wfo,zcode,gset)   

#    Save_Forecast('SanDiego',SanDiego_url_list,email=True)
#    Save_Forecast('LagunaBeach',LagunaBeach_url_list,email=True)
#    Save_Forecast('Riverside',Riverside_url_list,email=True)
#    Save_Forecast('Avalon',Avalon_url_list,email=True)
#    Save_Forecast('DanaPoint',Avalon_url_list,email=True)
#    Save_Forecast('JohnWayne',JohnWayne_url_list,email=True)
#    Save_Forecast('FoothillRanch',FoothillRanch_url_list,email=True)
#    Save_Forecast('ImperialBeach',ImperialBeach_url_list,email=True)
#    Save_Forecast('Irvine',Irvine_url_list,email=True)
#    Save_Forecast('Chollas',Chollas_url_list,email=True)
#    Save_Forecast('Lakeside',Lakeside_url_list,email=True)
#    Save_Forecast('RidingPark',RidingPark_url_list,email=True)
#    Save_Forecast('SR73',SR73_url_list,email=True)    
#    Save_Forecast('Fallbrook',Fallbrook_url_list,email=True)  
#    Save_Forecast('Pendleton',Pendleton_url_list,email=True)  
    Save_Forecast('SantaAnaFire',SantaAnaFire_url_list,email=True)  

   
    print 'Waiting for next forecast download interval...'

    print str(dt.datetime.now())
 

    #print 'Press any key to exit....'
    #input()
    print 'Forecasts saved.'



    
