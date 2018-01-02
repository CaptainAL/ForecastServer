# -*- coding: utf-8 -*-
"""
Created on Wed Nov 08 13:58:05 2017
@author: alex.messina
"""
if __name__=='__main__':
    import urllib
    import datetime as dt
    from docx import Document
    from docx.shared import Inches    
    import requests
    from email.MIMEMultipart import MIMEMultipart
    from email.MIMEText import MIMEText
    from email.MIMEImage import MIMEImage
    import smtplib
    
    ## Page with links: http://www.wrh.noaa.gov/sgx/obs/rainobs.php?wfo=sgx
    
    maindir = 'C:/Users/alex.messina/Documents/GitHub/ForecastServer/24HrRainTotalsSan Diego County/'
    ## Create Email
    msg = MIMEMultipart()
    ## Today's date
    now = str(dt.datetime.now())[:-13]+'00'
    ##  Create Document
    document = Document()
    ## Add Heading to email
    SUBJECT = 'NOAA 24Hr Rain Totals for San Diego County - '+now
    ## Add Heading to document
    document.add_heading(SUBJECT,level=1)
    ## Get Text forecast
    TEXT = requests.get('http://www.wrh.noaa.gov/sgx/data/hydro/LAXRRMSAN').content#.split('\n')
    ## Add Subject and Text to Email
    msg['SUBJECT']=SUBJECT
    content = 'Subject: {}\n\n{}'.format(SUBJECT, TEXT)
    msg.attach(MIMEText(content))
    ## Get and Save Picture
    pic_filename = maindir+'NOAA 24Hr Rain Totals for San Diego County_current.png'
    urllib.urlretrieve('http://www.cnrfc.noaa.gov/images/countyPrecipMaps/sandiego_24HRprecip.png', pic_filename)
    ## Add Picture to email
    msg.attach(MIMEImage(file(pic_filename,'rb').read()))
    ## Insert Picture into Document
    document.add_picture(pic_filename ,width=Inches(6)) ## add pic from filename defined above
    ## Add Text to Document (after Picture)
    paragraph = document.add_paragraph(TEXT)
    document.save(maindir+'NOAA 24Hr Rain Totals for San Diego County - '+now+'.docx')
    ## Send Email
    # to send
    mailer = smtplib.SMTP_SSL('smtp.googlemail.com:465')
    username, password = 'sebal.py', 'Mactec101'
    mailer.login(username,password)
    mailer.sendmail('sebal.py@gmail.com','atm1984@gmail.com', msg.as_string())
    mailer.close()
    print
    print 'Rain data saved uh successfurry . Have a nice day :)'
    print