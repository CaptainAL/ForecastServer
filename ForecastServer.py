# -*- coding: utf-8 -*-
"""
Created on Tue Oct 11 17:44:10 2016

@author: alex.messina
"""

import urllib
from docx import *
from docx.shared import Inches
import datetime as dt

from email.MIMEMultipart import MIMEMultipart
from email.MIMEText import MIMEText
from email.MIMEImage import MIMEImage
import smtplib


def Save_Forecast(site_name,url_list,maindir,email=False):
    print 'Looking up weather forecast for '+site_name+'....'
    print
    ##  Create Document
    document = Document()
    ## Create Email
    msg = MIMEMultipart()
    ## Today's date
    now = str(dt.datetime.now())[:-13]+'00'
    ## Add  Heading to Document
    header = 'NWS Forecast for '+site_name+' - '+now
    document.add_heading(header,level=1)
    ## Add Heading to email
    msg['SUBJECT']=header
    msg.attach(MIMEText(header))
    ##
    num_count = 1
    for url in url_list:
        print url
        ## Get and Save Picture
        pic_filename = maindir+'NWS '+site_name+' Forecast'+str(num_count)+'.png'
        urllib.urlretrieve(url, pic_filename )
        ## Insert Picture
        document.add_picture(pic_filename ,width=Inches(7)) ## add pic from filename defined above
        ## Add Picture to email
        msg.attach(MIMEImage(file(pic_filename,'rb').read()))
        ## count Up
        num_count += 1
    ## Save document
    print 'Saving forecasts...'
    document.save(maindir+'NWS Forecast for '+site_name+' - '+now+'.docx')
    if email==True:
        ## Send Email
        # to send
        mailer = smtplib.SMTP_SSL('smtp.googlemail.com:465')
        mailer.login(username,password)
        mailer.sendmail('sebal.py@gmail.com','atm1984@gmail.com', msg.as_string())
        mailer.close()
    print
    print 'Forecasts saved uh successfurry . Have a nice day :)'
    print
    return
    return

## San Diego
maindir = 'C:/Users/alex.messina/Documents/GitHub/ForecastServer/SanDiego/'
url_list = ['http://forecast.weather.gov/meteograms/Plotter.php?lat=32.7153&lon=-117.1573&wfo=SGX&zcode=CAZ043&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=0&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat=32.7153&lon=-117.1573&wfo=SGX&zcode=CAZ043&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=48&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat=32.7153&lon=-117.1573&wfo=SGX&zcode=CAZ043&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=96&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6']
Save_Forecast('San Diego',url_list,maindir,email=False)

## Laguna Beach
maindir = 'C:/Users/alex.messina/Documents/GitHub/ForecastServer/LagunaBeach/'
url_list=['http://forecast.weather.gov/meteograms/Plotter.php?lat=33.5422&lon=-117.7831&wfo=SGX&zcode=CAZ552&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=0&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat=33.5422&lon=-117.7831&wfo=SGX&zcode=CAZ552&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=48&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6', 'http://forecast.weather.gov/meteograms/Plotter.php?lat=33.5422&lon=-117.7831&wfo=SGX&zcode=CAZ552&gset=18&gdiff=3&unit=0&tinfo=PY8&ahour=96&pcmd=00000010100000000000000000000000000000000000000000000000000&lg=en&indu=1!1!1!&dd=&bw=&hrspan=48&pqpfhr=6&psnwhr=6']
Save_Forecast('LagunaBeach',url_list,maindir,email=False)

 